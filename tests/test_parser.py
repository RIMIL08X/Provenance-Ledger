"""Unit tests for the multi-format DataIngestionEngine."""

import io
import pandas as pd
import pytest
from docx import Document
from pypdf import PdfWriter

from src.parser.data_parser import DataIngestionEngine


def test_parse_csv():
    csv_content = b"name,age,score\nAlice,30,95.5\nBob,25,88.0\n"
    df, summary, data_hash = DataIngestionEngine.parse_file(csv_content, "test.csv")

    assert len(df) == 2
    assert "age" in df.columns
    assert summary["total_rows"] == 2
    assert len(data_hash) == 64


def test_parse_excel():
    df_in = pd.DataFrame({"product": ["A", "B"], "price": [10.0, 20.0]})
    buffer = io.BytesIO()
    df_in.to_excel(buffer, index=False)
    excel_bytes = buffer.getvalue()

    df, summary, data_hash = DataIngestionEngine.parse_file(excel_bytes, "test.xlsx")

    assert len(df) == 2
    assert "price" in df.columns
    assert summary["total_rows"] == 2
    assert len(data_hash) == 64


def test_parse_docx():
    doc = Document()
    doc.add_paragraph("First paragraph content")
    doc.add_paragraph("Second paragraph content")
    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()

    df, summary, data_hash = DataIngestionEngine.parse_file(docx_bytes, "document.docx")

    assert len(df) == 2
    assert "text" in df.columns
    assert len(data_hash) == 64


def test_load_stored_snapshot():
    csv_content = b"id,val\n1,100\n2,200\n"
    df, summary, data_hash = DataIngestionEngine.parse_file(csv_content, "snapshot_test.csv")

    loaded_df = DataIngestionEngine.load_snapshot(data_hash)
    assert loaded_df is not None
    assert len(loaded_df) == 2
    assert "val" in loaded_df.columns
