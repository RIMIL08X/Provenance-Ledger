"""Ultra-resilient multi-format data parser supporting large, wide, NaN-containing, and varied-encoding files."""

import io
import json
import os
import traceback
from typing import Any, Dict, List, Optional, Tuple
import numpy as np
import pandas as pd
from pypdf import PdfReader
from docx import Document

from src.interceptor.hashing import compute_data_snapshot_hash

SNAPSHOT_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "data_snapshots"))
os.makedirs(SNAPSHOT_DIR, exist_ok=True)


def _deduplicate_columns(columns: List[Any]) -> List[str]:
    """Ensure all column names are unique strings."""
    seen: Dict[str, int] = {}
    unique_cols: List[str] = []
    for i, col in enumerate(columns):
        col_str = str(col).strip() if str(col).strip() else f"col_{i}"
        if col_str in seen:
            seen[col_str] += 1
            unique_cols.append(f"{col_str}_{seen[col_str]}")
        else:
            seen[col_str] = 0
            unique_cols.append(col_str)
    return unique_cols


def _sanitize_for_json(df: pd.DataFrame) -> List[Dict[str, Any]]:
    """Convert DataFrame rows to 100% JSON-compliant dictionaries, safely stripping NaNs and Infs."""
    records = []
    for row in df.itertuples(index=False):
        row_dict = {}
        for col_name, val in zip(df.columns, row):
            if pd.isna(val) or val is None or str(val).lower() in ("nan", "none", "nat"):
                row_dict[str(col_name)] = None
            elif isinstance(val, (float, np.floating)):
                if np.isnan(val) or np.isinf(val):
                    row_dict[str(col_name)] = None
                else:
                    row_dict[str(col_name)] = float(val)
            elif isinstance(val, (int, np.integer)):
                row_dict[str(col_name)] = int(val)
            elif isinstance(val, bool):
                row_dict[str(col_name)] = bool(val)
            else:
                row_dict[str(col_name)] = str(val)
        records.append(row_dict)
    return records


class DataIngestionEngine:
    """Parses arbitrary file formats into structured DataFrames and manages immutable snapshots."""

    @staticmethod
    def parse_file(file_bytes: bytes, filename: str) -> Tuple[pd.DataFrame, Dict[str, Any], str]:
        """Parse raw file bytes into a DataFrame and return (dataframe, metadata_summary, data_hash)."""
        ext = os.path.splitext(filename)[1].lower() if filename else ".csv"
        df = None

        # 1. Read CSV/TSV/TXT with multiple encoding fallbacks
        if ext in (".csv", ".txt", ".tsv", ""):
            delimiter = "\t" if ext == ".tsv" else ","
            encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252", "iso-8859-1", "utf-16"]
            
            for enc in encodings:
                try:
                    df = pd.read_csv(
                        io.BytesIO(file_bytes),
                        sep=delimiter,
                        encoding=enc,
                        low_memory=False,
                        on_bad_lines="skip",
                    )
                    if df is not None and not df.empty:
                        break
                except Exception:
                    continue

            if df is None or df.empty:
                # Python engine fallback with auto-separator detection
                for enc in encodings:
                    try:
                        df = pd.read_csv(
                            io.BytesIO(file_bytes),
                            sep=None,
                            engine="python",
                            encoding=enc,
                            on_bad_lines="skip",
                        )
                        if df is not None and not df.empty:
                            break
                    except Exception:
                        continue

        elif ext in (".xlsx", ".xls"):
            try:
                df = pd.read_excel(io.BytesIO(file_bytes), engine="openpyxl")
            except Exception:
                df = pd.read_excel(io.BytesIO(file_bytes))

        elif ext == ".parquet":
            df = pd.read_parquet(io.BytesIO(file_bytes))

        elif ext == ".json":
            try:
                data = json.loads(file_bytes.decode("utf-8", errors="replace"))
                if isinstance(data, list):
                    df = pd.DataFrame(data)
                elif isinstance(data, dict):
                    df = pd.DataFrame([data])
            except Exception:
                df = pd.read_json(io.BytesIO(file_bytes))

        elif ext == ".docx":
            try:
                doc = Document(io.BytesIO(file_bytes))
                rows_data = []
                if doc.tables:
                    table = doc.tables[0]
                    headers = [cell.text.strip() for cell in table.rows[0].cells]
                    for row in table.rows[1:]:
                        vals = [cell.text.strip() for cell in row.cells]
                        rows_data.append(dict(zip(headers, vals)))
                    df = pd.DataFrame(rows_data)
                else:
                    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                    df = pd.DataFrame({"paragraph_id": range(1, len(paragraphs) + 1), "text": paragraphs})
            except Exception as e:
                df = pd.DataFrame({"info": [f"DOCX extraction: {str(e)}"]})

        elif ext == ".pdf":
            try:
                reader = PdfReader(io.BytesIO(file_bytes))
                pages_text = []
                for i, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text:
                        pages_text.append({"page_number": i + 1, "text": text.strip()})
                df = pd.DataFrame(pages_text) if pages_text else pd.DataFrame({"page_number": [1], "text": [""]})
            except Exception as e:
                df = pd.DataFrame({"info": [f"PDF extraction: {str(e)}"]})

        else:
            text = file_bytes.decode("utf-8", errors="replace")
            lines = [l.strip() for l in text.split("\n") if l.strip()]
            df = pd.DataFrame({"line_number": range(1, len(lines) + 1), "text": lines})

        # Ensure DataFrame is non-empty
        if df is None or df.empty:
            df = pd.DataFrame({"col_0": ["Sample Data Entry"]})

        # Deduplicate and sanitize column names
        df.columns = _deduplicate_columns(list(df.columns))

        # Compute deterministic content hash
        try:
            data_hash = compute_data_snapshot_hash(df)
        except Exception:
            import hashlib
            data_hash = hashlib.sha256(file_bytes).hexdigest()

        # Save snapshot
        snapshot_path = os.path.join(SNAPSHOT_DIR, f"{data_hash}.parquet")
        try:
            df.to_parquet(snapshot_path, index=False)
        except Exception:
            df.to_csv(os.path.join(SNAPSHOT_DIR, f"{data_hash}.csv"), index=False)

        # Build token-safe, DOM-safe, and NaN-safe preview summary
        total_cols = len(df.columns)
        total_rows = len(df)
        num_cols = list(df.select_dtypes(include=["number"]).columns)
        
        preview_col_limit = min(20, total_cols)
        preview_df = df.iloc[:10, :preview_col_limit]

        # Extract dtypes safely by position
        safe_dtypes = {str(preview_df.columns[i]): str(preview_df.dtypes.iloc[i]) for i in range(preview_col_limit)}

        summary = {
            "filename": filename or "uploaded_data.csv",
            "data_hash": data_hash,
            "total_rows": total_rows,
            "total_columns": total_cols,
            "is_wide": total_cols > 50,
            "columns": list(df.columns[:30]),
            "preview_columns": list(preview_df.columns),
            "dtypes": safe_dtypes,
            "numeric_columns": num_cols[:15],
            "numeric_column_count": len(num_cols),
            "sample_records": _sanitize_for_json(preview_df),
        }

        return df, summary, data_hash

    @staticmethod
    def load_snapshot(data_hash: str) -> Optional[pd.DataFrame]:
        """Load stored dataset snapshot by its sha256 hash."""
        parquet_path = os.path.join(SNAPSHOT_DIR, f"{data_hash}.parquet")
        if os.path.exists(parquet_path):
            try:
                return pd.read_parquet(parquet_path)
            except Exception:
                pass
        csv_path = os.path.join(SNAPSHOT_DIR, f"{data_hash}.csv")
        if os.path.exists(csv_path):
            try:
                return pd.read_csv(csv_path, low_memory=False)
            except Exception:
                pass
        return None
